#!/usr/bin/env python3
"""
parse_bolton.py  -- Bolton side of the Bolton-vs-AntCat catalogue diff.

Reads the unzipped AA_CATALOGUE folder of .docx files and writes:
  bolton_species.csv     one row per species-group headword
  bolton_genera.csv      one row per genus-group headword
  bolton_references.csv  one row per bibliographic entry

Parsing relies on Bolton's consistent run-level formatting:
  * species headword  = first run BOLD+ITALIC, lowercase epithet  ("brevicornis. Acanthognathus ...")
  * genus header      = leading ALL-CAPS token;
                          BOLD+ITALIC  -> valid genus/subgenus  (bracket = classification)
                          ITALIC-only  -> synonym/homonym/etc.  (bracket = status)
  * reference         = plain paragraph in a REFERENCES file, "Author, Init. Year. Title ..."

Matching key for species = (current_genus, gstem(headword_epithet)).  The current
genus is the file's ALL-CAPS section header (fallback: filename).  We do NOT try to
fully parse the original-combination string -- the headword epithet is the terminal
epithet, which is all the (genus, terminal) key needs.
"""
import csv, sys, re, os, json, unicodedata
from docx import Document
from parse_antcat import gstem, norm_genus

YEAR = re.compile(r'\b(1[6789]\d\d|20\d\d)([a-z]?)\b')

# Detecting references where Bolton omitted the publication year, so the first
# 4-digit number in the line is a taxonomic authority inside the TITLE (e.g.
# "...genus Syscia Roger, 1816...") rather than the year of publication. The
# tell is that the text the parser would take as the "author" then contains a
# run of title words. Folding accents (incl. the few letters NFKD leaves alone,
# like O-slash) keeps Central/European surnames from fragmenting into false
# title words, and consortium-style author tokens are whitelisted, so this fires
# only on the genuine case.
_XMAP = str.maketrans({'Ø': 'O', 'ø': 'o', 'Đ': 'D', 'đ': 'd', 'ł': 'l',
                       'Ł': 'L', 'ı': 'i', 'ð': 'd', 'Ð': 'D', 'þ': 't'})
_NAME_WORDS = {'de', 'van', 'von', 'da', 'do', 'dos', 'del', 'della', 'le', 'la',
               'den', 'der', 'ten', 'ter', 'di', 'du', 'af', 'zu', 'y', 'e', 'in',
               'das', 'dem', 'ed', 'jr', 'sr', 'und', 'el', 'consortium', 'plus',
               'more', 'authors', 'team', 'group', 'network'}

def _looks_like_title(author):
    """True if 'author' holds >=2 consecutive lowercase title words (=> the year
    the parser found is a title-internal authority date, not the pub year)."""
    a = author.translate(_XMAP)
    a = unicodedata.normalize('NFKD', a)
    a = ''.join(c for c in a if not unicodedata.combining(c))
    a = re.sub(r'\([^)]*\)', ' ', a)
    run = best = 0
    for w in re.findall(r"[A-Za-z'-]+", a):
        if w == w.lower() and len(w) > 1 and w.lower() not in _NAME_WORDS:
            run += 1; best = max(best, run)
        else:
            run = 0
    return best >= 2

def first_run(p):
    return next((r for r in p.runs if r.text.strip()), None)

def leading_italic(p):
    s = ''
    for r in p.runs:
        if r.italic:
            s += r.text
        else:
            break
    return s

# ---------------------------------------------------------------- species
def parse_species_file(path):
    d = Document(path)
    fname_genus = re.sub(r'\.docx?$', '', os.path.basename(path))
    fname_genus = re.sub(r'^CAT-SPECIES\s+', '', fname_genus).strip()
    # strip trailing alphabetic-range suffixes on split files: "CAMPONOTUS a-b" -> "CAMPONOTUS"
    fname_genus = re.sub(r'\s+[a-zA-Z](\s*-\s*[a-zA-Z])?$', '', fname_genus).strip().title()
    cur_genus = fname_genus
    out = []
    cur_lines = None                      # accumulates the current headword's full block
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        fr = first_run(p)
        if not fr:
            continue
        # section genus header: a single ALL-CAPS token, bold+italic
        if re.fullmatch(r'\*?[A-Z][A-Z\-]+', t) and fr.bold and fr.italic:
            cur_genus = t.replace('*', '').strip().title()
            continue
        ft = fr.text.strip()
        if fr.bold and fr.italic and re.match(r'^\*?[a-z]', ft):
            # headword epithet comes from the TEXT (runs sometimes split the word),
            # the bold+italic first run is only the signal that this is a headword.
            hm = re.match(r'^\*?([a-z][a-z\-]+)\.', t)
            if not hm:
                if cur_lines is not None:          # not a real headword -> continuation
                    cur_lines.append(t)
                continue
            if out and cur_lines is not None:      # finalise the previous block
                out[-1]['block_text'] = '\n'.join(cur_lines)
            epithet = hm.group(1)
            fossil = t.startswith('*') or ft.startswith('*')
            # original-combination string + rank, parsed from TEXT (after the headword)
            after = t[hm.end():].strip()
            ym = YEAR.search(after)
            year = (ym.group(1) + ym.group(2)) if ym else ''
            namepart = after[:ym.start()].strip() if ym else after
            toks = namepart.split()
            name_toks = [toks[0]] if toks else []          # genus
            author = ''
            for w in toks[1:]:
                w0 = w.strip('(),.')
                if w0 in ('var', 'subsp', 'r', 'st', 'n', 'f', 'nr', 'in', 'sp'):
                    continue
                if w.startswith('('):                       # subgenus in parens
                    continue
                if re.match(r'^[a-z]', w0):
                    name_toks.append(w0)
                else:
                    author = w0                             # first author surname
                    break
            orig = ' '.join(name_toks).strip()
            n_ep = max(0, len(name_toks) - 1)
            rank = 'species' if n_ep <= 1 else 'subspecies' if n_ep == 2 else 'infrasub'
            out.append(dict(genus=cur_genus, epithet=epithet,
                            gstem=gstem(epithet), rank=rank, fossil=fossil,
                            year=year, original_combination=orig, author=author,
                            headword_text=t[:240], block_text=t,
                            source_file=os.path.basename(path)))
            cur_lines = [t]                                  # start the new block
        else:
            if cur_lines is not None:                        # sub-line of current headword
                cur_lines.append(t)
    if out and cur_lines is not None:                        # finalise the last block
        out[-1]['block_text'] = '\n'.join(cur_lines)
    return out

# ---------------------------------------------------------------- genera
def parse_genus_file(path):
    d = Document(path)
    out = []
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        fr = first_run(p)
        if not fr:
            continue
        m = re.match(r'^([A-Z][A-Z\-]{2,})\b', t)
        if not m:
            continue
        name = m.group(1)
        if name in ('BARRY', 'FORMICIDAE'):
            continue
        br = re.search(r'\[([^\]]*)\]', t)
        bracket = br.group(1).strip() if br else ''
        colon = ''
        cm = re.match(r'^[A-Z][A-Z\-]+:\s*(.*)$', t)
        if cm:
            colon = cm.group(1).strip()[:80]
        if fr.bold and fr.italic:
            status = 'valid'
        elif fr.italic:
            status = 'invalid'        # synonym / homonym / misspelling / nomen nudum
        else:
            continue
        out.append(dict(name=name.title(), key=norm_genus(name), status=status,
                        classification=bracket or colon, header_text=t[:200],
                        source_file=os.path.basename(path)))
    return out

# ---------------------------------------------------------------- references
def parse_reference_file(path):
    d = Document(path)
    out = []
    for p in d.paragraphs:
        t = p.text.strip()
        if not t or t.startswith('[') or t.startswith('BARRY') or t.upper() == t and len(t) < 40:
            continue
        # a reference line starts with an author surname and contains "Year." early
        if not re.match(r'^[A-ZÀ-Þ]', t):
            continue
        ym = YEAR.search(t)
        if not ym:
            continue
        year = ym.group(1) + ym.group(2)
        author = t[:ym.start()].strip().rstrip(',').strip()
        if len(author) < 2:
            continue
        # Bolton omitted the year and the parser caught a title authority date:
        # blank the year + author_key so it can't form a (wrong) match key and
        # drops out of the diff instead of showing as a phantom "missing" ref.
        malformed = _looks_like_title(author)
        if malformed:
            year, author_key, key = '', '', ''
        else:
            author_key = re.sub(r'[^a-z]', '', author.lower())
            key = f'{author_key}|{year}'
        out.append(dict(author=author, year=year,
                        author_key=author_key, key=key, malformed_year=malformed,
                        reference_text=t[:300], source_file=os.path.basename(path)))
    return out

# ---------------------------------------------------------------- driver
def build(folder, outdir):
    os.makedirs(outdir, exist_ok=True)
    files = sorted(os.listdir(folder))
    sp, gn, rf = [], [], []
    for fn in files:
        full = os.path.join(folder, fn)
        if fn.startswith('CAT-SPECIES'):
            sp += parse_species_file(full)
        elif fn.startswith('CAT-GENUS GROUP TAXA'):
            gn += parse_genus_file(full)
        elif re.match(r'^SUB\.\s*\d+\s*REFERENCES', fn):
            rf += parse_reference_file(full)

    def dump(rows, name):
        path = os.path.join(outdir, name)
        with open(path, 'w', newline='', encoding='utf-8') as f:
            w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
            w.writeheader(); w.writerows(rows)
        return path

    # full headword blocks for the content-level diff (diff_history.py)
    with open(os.path.join(outdir, 'bolton_species_blocks.jsonl'), 'w', encoding='utf-8') as f:
        for r in sp:
            rec = {k: r.get(k) for k in ('genus', 'epithet', 'gstem', 'rank', 'fossil',
                                         'year', 'original_combination', 'author', 'block_text')}
            f.write(json.dumps(rec, ensure_ascii=False) + '\n')
    for r in sp:                              # keep block_text out of the flat CSV
        r.pop('block_text', None)

    dump(sp, 'bolton_species.csv')
    dump(gn, 'bolton_genera.csv')
    dump(rf, 'bolton_references.csv')

    print(f"Bolton species-group headwords: {len(sp):>6}  "
          f"(species={sum(r['rank']=='species' for r in sp)}, "
          f"subspecies={sum(r['rank']=='subspecies' for r in sp)}, "
          f"infrasub={sum(r['rank']=='infrasub' for r in sp)})")
    print(f"Bolton genus-group headwords:   {len(gn):>6}  "
          f"(valid={sum(r['status']=='valid' for r in gn)}, "
          f"invalid={sum(r['status']=='invalid' for r in gn)})")
    print(f"Bolton reference entries:       {len(rf):>6}")
    return sp, gn, rf

if __name__ == '__main__':
    folder = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else '.'
    build(folder, out)
